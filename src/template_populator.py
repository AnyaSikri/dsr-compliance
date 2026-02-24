"""Template populator: assemble a filled markdown document and .docx from parsed template sections.

Walks parsed template sections, resolves source references using ib_resolver,
synthesizes report-ready prose via LLM, and produces a single filled markdown
document plus a .docx conversion.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document

from src.ib_resolver import resolve_sources
from src.models import TemplateSection
from src.utils import ensure_dir, logger

if TYPE_CHECKING:
    from src.openai_client import LLMClient


# ---------------------------------------------------------------------------
# LLM synthesis prompt
# ---------------------------------------------------------------------------

SYNTHESIS_SYSTEM = """\
You are a regulatory medical writer producing a Drug Safety Report (DSR). \
Given source material extracted from reference documents (Investigator's \
Brochure, PBRER, literature databases), write the content for the specified \
section of the DSR.

RULES — follow these exactly:
1. Write in formal regulatory prose suitable for a DSR submission to health \
authorities.
2. Use ONLY the information provided in the source material. Do not add \
facts, statistics, or claims that are not supported by the provided sources.
3. Preserve all specific data points exactly: numbers, percentages, dates, \
study names, MedDRA preferred terms, patient counts, p-values.
4. Format any tabular data as markdown tables.
5. Do NOT include the template instructions in your output — they are \
guidance for you, not report content.
6. Do NOT add disclaimers, meta-commentary, or notes about your writing \
process.
7. If the source material is clearly insufficient for a complete section, \
write what you can from the available data and add a single line: \
"[ADDITIONAL DATA NEEDED: brief description of what is missing]"
8. Use section-appropriate structure: bullet lists for indications, \
prose paragraphs for background/discussion, structured summaries for \
results with quantitative data.
9. Do NOT repeat the section heading — it is already added by the system.\
"""


def _build_synthesis_prompt(
    section: TemplateSection,
    source_contents: list[tuple[str, str]],
) -> str:
    """Build the user prompt for LLM synthesis.

    Args:
        section: The template section being populated.
        source_contents: List of (source_label, content) tuples from resolved sources.
    """
    parts: list[str] = []
    parts.append(f"SECTION: {section.section_id} — {section.title}")
    parts.append("")

    if section.body:
        parts.append("TEMPLATE INSTRUCTIONS (for your guidance, do NOT include these in the output):")
        parts.append(section.body)
        parts.append("")

    if source_contents:
        parts.append("SOURCE MATERIAL:")
        for label, content in source_contents:
            parts.append(f"\n--- {label} ---")
            # Truncate very long sources to stay within token limits
            if len(content) > 12000:
                parts.append(content[:12000])
                parts.append("[... content truncated for length ...]")
            else:
                parts.append(content)
        parts.append("")
    else:
        parts.append("SOURCE MATERIAL: None available.")
        parts.append(
            "Write a structured placeholder noting what data this section "
            "requires and from which sources, based on the template instructions above."
        )

    return "\n".join(parts)


def _heading_level(section_id: str) -> int:
    """Determine the markdown heading level from a section_id's depth.

    - "1" -> ## (level 2)
    - "2.1" -> ### (level 3)
    - "2.1.1" -> #### (level 4)
    - Non-numeric ids (e.g. "Executive Summary") -> ## (level 2)
    - Cap at 6
    """
    parts = section_id.strip().split(".")
    try:
        for p in parts:
            int(p)
        depth = len(parts)
        level = depth + 1
        return min(level, 6)
    except ValueError:
        return 2


def assemble_markdown(
    template_sections: list[TemplateSection],
    ib_index: dict[str, str],
    llm: LLMClient | None = None,
    dry_run: bool = False,
    pbrer_index: dict[str, str] | None = None,
    literature_results: dict[str, str] | None = None,
) -> str:
    """Build a single markdown document from template sections and resolved content.

    When *llm* is provided and *dry_run* is False, each section's resolved
    source material is sent through the LLM for synthesis into report-ready
    prose.  In dry-run mode or when no LLM is available, the legacy behavior
    (raw source paste / template body) is preserved.
    """
    lines: list[str] = ["# Filled Signal Assessment Report\n"]
    use_synthesis = llm is not None and not dry_run

    # Build lookup structures for dedup: skip children that have no
    # required_sources when their parent section is already in the list
    # (the parent's synthesized body already covers them).
    all_ids = {s.section_id for s in template_sections}
    sections_with_sources = {
        s.section_id for s in template_sections if s.required_sources
    }

    def _parent_id(section_id: str) -> str | None:
        """Return the parent section ID, e.g. '3.1' for '3.1.2'."""
        parts = section_id.rsplit(".", 1)
        return parts[0] if len(parts) > 1 else None

    def _should_skip_child(section_id: str) -> bool:
        """Skip a child section if it has no sources and its parent exists."""
        pid = _parent_id(section_id)
        return pid is not None and pid in all_ids and section_id not in sections_with_sources

    # Track section content for Executive Summary second pass
    section_contents: dict[str, str] = {}

    for section in template_sections:
        level = _heading_level(section.section_id)
        hashes = "#" * level
        lines.append(f"{hashes} {section.section_id} {section.title}\n")

        # Skip children without their own sources when the parent section
        # is in the list — the parent's body already covers them.
        if _should_skip_child(section.section_id):
            continue

        # Check if this is an Executive Summary subsection (1.x) — defer
        # to the second pass so it can summarize the completed report.
        is_exec_summary = (
            section.section_id.startswith("1.")
            and not section.required_sources
            and use_synthesis
        )
        if is_exec_summary:
            # Placeholder marker; will be replaced in second pass
            lines.append(f"{{{{EXEC_SUMMARY_{section.section_id}}}}}\n")
            continue

        section_text = ""

        if not section.required_sources:
            # No sources referenced in template
            if use_synthesis and section.body:
                prompt = _build_synthesis_prompt(section, [])
                try:
                    section_text = llm.call(
                        system_prompt=SYNTHESIS_SYSTEM,
                        user_prompt=prompt,
                        json_mode=False,
                        label=f"synth_{section.section_id}",
                    ).strip()
                except Exception as e:
                    logger.warning(
                        "Synthesis failed for %s: %s — falling back to template body",
                        section.section_id, e,
                    )
                    section_text = section.body
            elif section.body:
                section_text = section.body
        else:
            # Resolve sources
            resolved = resolve_sources(
                section.required_sources,
                ib_index,
                pbrer_index=pbrer_index,
                literature_results=literature_results,
            )

            if use_synthesis:
                source_contents: list[tuple[str, str]] = []
                for rs in resolved:
                    source_contents.append((rs.original_ref, rs.content))

                prompt = _build_synthesis_prompt(section, source_contents)
                try:
                    section_text = llm.call(
                        system_prompt=SYNTHESIS_SYSTEM,
                        user_prompt=prompt,
                        json_mode=False,
                        label=f"synth_{section.section_id}",
                    ).strip()
                except Exception as e:
                    logger.warning(
                        "Synthesis failed for %s: %s — falling back to raw sources",
                        section.section_id, e,
                    )
                    raw_lines: list[str] = []
                    _append_raw_sources(raw_lines, resolved)
                    section_text = "\n".join(raw_lines)
            else:
                raw_lines = []
                _append_raw_sources(raw_lines, resolved)
                section_text = "\n".join(raw_lines)

        if section_text:
            lines.append(f"{section_text}\n")
            section_contents[section.section_id] = section_text

    # --- Second pass: fill Executive Summary from completed report ---
    if use_synthesis:
        _fill_executive_summary(lines, template_sections, section_contents, llm)

    return "\n".join(lines)


EXEC_SUMMARY_SYSTEM = """\
You are a regulatory medical writer producing the Executive Summary section \
of a Drug Safety Report (DSR). You are given the completed content from the \
main body sections of the report. Write a concise executive summary \
subsection that synthesizes the key points from the report body.

RULES:
1. Write in formal regulatory prose.
2. Base your summary ONLY on the report content provided — do not add \
information not present in the source material.
3. Be concise — each subsection should be 2-5 sentences.
4. Preserve all specific data points: numbers, drug names, study names.
5. Do NOT repeat the section heading.
6. If there is genuinely insufficient data in the report body for a \
particular summary subsection, write what you can and add: \
"[ADDITIONAL DATA NEEDED: brief description]"\
"""


def _fill_executive_summary(
    lines: list[str],
    template_sections: list[TemplateSection],
    section_contents: dict[str, str],
    llm: LLMClient,
) -> None:
    """Replace Executive Summary placeholders with content from report body.

    Feeds the completed report sections (3-7) into the LLM to generate
    each Executive Summary subsection (1.1-1.6).
    """
    # Collect report body for context (sections 2+, truncated to fit)
    body_parts: list[str] = []
    for sid in sorted(section_contents.keys()):
        if not sid.startswith("1"):
            body_parts.append(f"--- Section {sid} ---\n{section_contents[sid]}")
    report_body = "\n\n".join(body_parts)
    # Truncate to ~30K chars to stay within token limits
    if len(report_body) > 30000:
        report_body = report_body[:30000] + "\n[... truncated ...]"

    if not report_body.strip():
        logger.warning("No report body content available for Executive Summary synthesis")
        return

    exec_sections = [
        s for s in template_sections
        if s.section_id.startswith("1.") and not s.required_sources
    ]

    for section in exec_sections:
        marker = f"{{{{EXEC_SUMMARY_{section.section_id}}}}}"
        # Check if marker exists in lines
        found = False
        for i, line in enumerate(lines):
            if marker in line:
                found = True
                prompt = (
                    f"SECTION: {section.section_id} — {section.title}\n\n"
                    f"TEMPLATE GUIDANCE: {section.body or 'Summarize relevant findings.'}\n\n"
                    f"COMPLETED REPORT BODY:\n{report_body}"
                )
                try:
                    content = llm.call(
                        system_prompt=EXEC_SUMMARY_SYSTEM,
                        user_prompt=prompt,
                        json_mode=False,
                        label=f"exec_{section.section_id}",
                    ).strip()
                    lines[i] = f"{content}\n"
                    logger.info("Filled Executive Summary %s from report body", section.section_id)
                except Exception as e:
                    logger.warning(
                        "Executive Summary synthesis failed for %s: %s",
                        section.section_id, e,
                    )
                    lines[i] = (
                        f"[ADDITIONAL DATA NEEDED: Executive summary for "
                        f"{section.title} — synthesis failed.]\n"
                    )
                break
        if not found:
            logger.debug("No placeholder found for exec summary %s", section.section_id)


def _append_raw_sources(lines: list[str], resolved: list) -> None:
    """Append resolved sources in the legacy raw-paste format."""
    if len(resolved) == 1:
        rs = resolved[0]
        lines.append(f"*Source: {rs.original_ref}*\n")
        lines.append(f"{rs.content}\n")
    else:
        for rs in resolved:
            lines.append(f"### From {rs.original_ref}\n")
            lines.append(f"{rs.content}\n")


def _add_field_code(run, field_code: str) -> None:
    """Insert a Word field code (e.g. PAGE, NUMPAGES, TOC) into a run."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def _setup_document(doc: Document) -> None:
    """Configure document layout: margins, fonts, headers, footers, TOC."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # --- Page layout ---
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True

    # --- Default font ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # --- Heading styles ---
    for level in range(1, 5):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = "Calibri"
            h_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # --- Title page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run("SIGNAL ASSESSMENT REPORT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Drug Safety Report")
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    confidential.paragraph_format.space_before = Pt(48)
    run = confidential.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

    # --- Table of Contents page ---
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()
    _add_field_code(toc_run, 'TOC \\o "1-3" \\h \\z \\u')

    doc.add_page_break()

    # --- Footer with page numbers (pages after title) ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Page ")
    page_run = footer_para.add_run()
    _add_field_code(page_run, "PAGE")
    footer_para.add_run(" of ")
    total_run = footer_para.add_run()
    _add_field_code(total_run, "NUMPAGES")

    # --- Header (pages after title) ---
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("Signal Assessment Report — Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True


def _add_rich_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with inline markdown formatting (bold, italic).

    Handles **bold**, *italic*, and mixed formatting within a single
    paragraph.
    """
    p = doc.add_paragraph()
    # Split on bold/italic markers and create runs
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def _add_markdown_table(doc: Document, lines: list[str], start_idx: int) -> int:
    """Parse a markdown table starting at start_idx and add it to the doc.

    Returns the index of the first line after the table.
    """
    from docx.shared import Pt, RGBColor
    import docx.oxml

    table_lines = []
    idx = start_idx
    while idx < len(lines) and "|" in lines[idx]:
        table_lines.append(lines[idx].strip())
        idx += 1

    if len(table_lines) < 2:
        return start_idx

    # Parse header
    header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]

    # Skip separator line (e.g. |---|---|)
    data_start = 1
    if data_start < len(table_lines) and re.match(r"^[\|\s\-:]+$", table_lines[data_start]):
        data_start = 2

    # Parse data rows
    data_rows = []
    for line in table_lines[data_start:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            data_rows.append(cells)

    if not header_cells:
        return start_idx

    # Create table
    num_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"

    # Header row
    for i, cell_text in enumerate(header_cells):
        if i < num_cols:
            cell = table.rows[0].cells[i]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
            # Header shading
            shading = docx.oxml.parse_xml(
                f'<w:shd {docx.oxml.ns.nsdecls("w")} w:fill="1F3A5F"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows with alternating shading
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                # Alternating row color
                if row_idx % 2 == 0:
                    shading = docx.oxml.parse_xml(
                        f'<w:shd {docx.oxml.ns.nsdecls("w")} w:fill="F2F2F2"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

    # Add spacing after table
    doc.add_paragraph()
    return idx


def _markdown_to_docx(md_content: str, output_path: Path) -> None:
    """Convert markdown text to a professional .docx file.

    Produces a document with:
    - Title page with report name and confidentiality notice
    - Auto-updating Table of Contents
    - Headers and footers with page numbers
    - Proper heading styles (Calibri, navy blue)
    - Formatted tables with header shading and alternating rows
    - Inline bold/italic formatting
    - Highlighted placeholder sections for missing data
    """
    from docx.shared import Pt, RGBColor

    doc = Document()
    _setup_document(doc)

    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        # Skip the top-level title (already on the title page)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        # Heading lines
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            heading_text = heading_match.group(2)
            level = len(heading_match.group(1))
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # Markdown table
        if "|" in stripped and stripped.startswith("|"):
            new_i = _add_markdown_table(doc, lines, i)
            if new_i > i:
                i = new_i
                continue

        # Bullet list items
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            _add_rich_paragraph(doc, bullet_match.group(1))
            doc.paragraphs[-1].style = "List Bullet"
            i += 1
            continue

        # Placeholder lines — highlighted in yellow-ish with bold
        if (
            stripped.startswith("[MANUAL INPUT REQUIRED:")
            or stripped.startswith("[CONTENT NOT FOUND:")
            or stripped.startswith("[ADDITIONAL DATA NEEDED:")
        ):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            run.font.size = Pt(10)
            i += 1
            continue

        # Normal paragraph with inline formatting
        _add_rich_paragraph(doc, stripped)
        i += 1

    doc.save(str(output_path))


def write_filled_template(
    template_sections: list[TemplateSection],
    ib_index: dict[str, str],
    output_dir: str | Path,
    llm: LLMClient | None = None,
    dry_run: bool = False,
    pbrer_index: dict[str, str] | None = None,
    literature_results: dict[str, str] | None = None,
) -> dict[str, Path]:
    """Write filled_template.md and filled_template.docx, return their paths.

    When *llm* is provided, source material is synthesized into report-ready
    prose.  Pass ``dry_run=True`` to skip synthesis and use legacy raw-paste
    behavior.

    Returns:
        ``{"md": Path(...), "docx": Path(...)}``
    """
    output_dir = Path(output_dir)
    ensure_dir(output_dir)

    md_content = assemble_markdown(
        template_sections, ib_index,
        llm=llm,
        dry_run=dry_run,
        pbrer_index=pbrer_index,
        literature_results=literature_results,
    )

    md_path = output_dir / "filled_template.md"
    md_path.write_text(md_content, encoding="utf-8")
    logger.info("Wrote filled markdown template to %s", md_path)

    docx_path = output_dir / "filled_template.docx"
    _markdown_to_docx(md_content, docx_path)
    logger.info("Wrote filled DOCX template to %s", docx_path)

    return {"md": md_path, "docx": docx_path}
