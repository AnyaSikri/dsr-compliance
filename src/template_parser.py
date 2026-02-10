"""Parse a regulatory template (.txt or .docx) into TemplateSection objects.

Uses two API calls:
  1. Section identification (structure extraction)
  2. Verbatim source extraction per section

Supports .docx templates with mapping table parsing and IGNORE section
detection. Post-extraction validation ensures every extracted source
appears as an exact substring in the template text. Results are cached
to data/intermediate/ to avoid redundant API calls on re-runs.
"""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from .config import Config
from .models import MappingTableEntry, TemplateSection
from .openai_client import LLMClient
from .utils import ensure_dir, logger


def _template_hash(template_path: Path) -> str:
    """Hash template file bytes for caching (works for both .txt and .docx)."""
    return hashlib.sha256(template_path.read_bytes()).hexdigest()[:16]


# ---------------------------------------------------------------------------
# .docx reading support
# ---------------------------------------------------------------------------


def _read_template_content(
    template_path: Path,
) -> tuple[str, list[list[list[str]]]]:
    """Read template content, handling both .txt and .docx formats.

    Returns (text_content, tables) where tables is a list of tables
    extracted from .docx (empty list for .txt files).
    """
    if template_path.suffix.lower() in (".docx", ".doc"):
        from docx import Document as DocxDocument

        doc = DocxDocument(str(template_path))
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)

        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                tables.append(rows)

        return text, tables

    return template_path.read_text(encoding="utf-8"), []


def read_template_text(template_path: Path) -> str:
    """Read template text content, supporting both .txt and .docx.

    Public convenience function for use by cli.py and validators.
    """
    text, _ = _read_template_content(template_path)
    return text


# ---------------------------------------------------------------------------
# Mapping table parsing
# ---------------------------------------------------------------------------

_HEADER_KEYWORDS = {"section", "content", "source", "reference", "document"}

_SOURCE_SPLIT_RE = re.compile(r"\s+OR\s+|[;,]\s*", re.IGNORECASE)


def _split_source_refs(text: str) -> list[str]:
    """Split a source reference string into individual refs."""
    if not text.strip():
        return []
    parts = _SOURCE_SPLIT_RE.split(text)
    return [p.strip() for p in parts if p.strip()]


def _find_column(header: list[str], keywords: list[str]) -> int | None:
    """Find the column index whose header contains one of the keywords."""
    for i, cell in enumerate(header):
        cell_lower = cell.lower().strip()
        for kw in keywords:
            if kw in cell_lower:
                return i
    return None


def _parse_mapping_table(
    tables: list[list[list[str]]],
) -> list[MappingTableEntry]:
    """Parse the mapping table from the template's extracted tables.

    Looks for a table with headers containing keywords like 'section',
    'source', 'content'. Returns MappingTableEntry objects for each row.
    """
    if not tables:
        return []

    for table in tables:
        if not table or len(table) < 2:
            continue
        header_text = " ".join(cell.lower() for cell in table[0])
        if any(kw in header_text for kw in _HEADER_KEYWORDS):
            return _parse_table_rows(table)

    return []


def _parse_table_rows(table: list[list[str]]) -> list[MappingTableEntry]:
    """Parse rows of the identified mapping table."""
    header = table[0]
    section_col = _find_column(header, ["section", "dsr section", "dsr"])
    source_col = _find_column(header, ["source", "reference"])

    if section_col is None:
        section_col = 0
    if source_col is None:
        source_col = min(2, len(header) - 1)

    entries: list[MappingTableEntry] = []
    for row in table[1:]:
        if not any(cell.strip() for cell in row):
            continue

        section_text = row[section_col].strip() if section_col < len(row) else ""
        source_text = row[source_col].strip() if source_col < len(row) else ""

        # Extract section number from the section text
        m = re.match(r"(\d+(?:\.\d+)*)", section_text)
        if not m:
            continue

        section_id = m.group(1)
        title = section_text[m.end():].strip().lstrip("-").strip()
        source_refs = _split_source_refs(source_text)

        entries.append(
            MappingTableEntry(
                dsr_section_id=section_id,
                dsr_title=title,
                source_refs=source_refs,
            )
        )

    return entries


def get_mapping_entries(template_path: Path) -> list[MappingTableEntry]:
    """Extract mapping table entries from a template file.

    Public function for use by cli.py to pass mapping entries to the
    section mapper.
    """
    _, tables = _read_template_content(template_path)
    return _parse_mapping_table(tables)


# ---------------------------------------------------------------------------
# IGNORE section handling
# ---------------------------------------------------------------------------

_IGNORE_PATTERNS = [
    re.compile(r"\bignore\b", re.IGNORECASE),
    re.compile(r"previous\s+template\s+version", re.IGNORECASE),
    re.compile(r"do\s+not\s+use", re.IGNORECASE),
]


def _mark_ignore_sections(sections: list[TemplateSection]) -> None:
    """Mark sections that should be excluded from processing.

    When an IGNORE marker is found, that section and all following
    sections are marked with ignore=True.
    """
    ignore_from: int | None = None
    for i, section in enumerate(sections):
        full_text = f"{section.section_id} {section.title} {section.body[:200]}"
        for pattern in _IGNORE_PATTERNS:
            if pattern.search(full_text):
                ignore_from = i
                break
        if ignore_from is not None:
            break

    if ignore_from is not None:
        logger.info(
            "Marking sections from index %d onwards as IGNORE (%d sections)",
            ignore_from,
            len(sections) - ignore_from,
        )
        for section in sections[ignore_from:]:
            section.ignore = True


def _apply_mapping_table(
    sections: list[TemplateSection],
    mapping_entries: list[MappingTableEntry],
) -> None:
    """Enrich TemplateSection.mapping_table_sources from mapping table entries."""
    from .ib_resolver import classify_source

    entry_by_id: dict[str, MappingTableEntry] = {
        e.dsr_section_id: e for e in mapping_entries
    }

    for section in sections:
        entry = entry_by_id.get(section.section_id)
        if entry is None:
            continue

        # Classify each source ref and group by type
        grouped: dict[str, list[str]] = {}
        for ref in entry.source_refs:
            source_type, _ = classify_source(ref)
            grouped.setdefault(source_type, []).append(ref)

        section.mapping_table_sources = grouped

        # Merge into required_sources if not already present
        for ref in entry.source_refs:
            if ref not in section.required_sources:
                section.required_sources.append(ref)


def _load_cache(cache_path: Path) -> list[TemplateSection] | None:
    if not cache_path.exists():
        return None
    try:
        data = json.loads(cache_path.read_text(encoding="utf-8"))
        return [TemplateSection(**s) for s in data["sections"]]
    except Exception:
        return None


def _save_cache(cache_path: Path, sections: list[TemplateSection], text_hash: str) -> None:
    data = {
        "text_hash": text_hash,
        "sections": [s.model_dump() for s in sections],
    }
    cache_path.write_text(json.dumps(data, indent=2), encoding="utf-8")


# --------------------------------------------------------------------------
# API prompts
# --------------------------------------------------------------------------

SECTION_ID_SYSTEM = """\
You are a regulatory document parser. Given the full text of a regulatory \
template, identify every discrete section. Return a JSON object with key \
"sections" containing a list of objects, each with:
  - "section_id": the section number or identifier (e.g. "2.1.1", \
"Executive Summary", "Appendices")
  - "title": the section heading
  - "body": the full text belonging to that section (everything until \
the next section heading)

Include ALL sections, even those with minimal body text. Preserve the \
exact wording of headings and body text. Do not infer, merge, or skip \
sections. If a section has sub-sections, list each sub-section as its \
own entry AND include the parent section.\
"""

SOURCE_EXTRACT_SYSTEM = """\
You are a regulatory compliance specialist. For each template section \
provided, extract ONLY the sources that are EXPLICITLY named in the \
template text.

RULES — follow these exactly:
1. Copy each source reference VERBATIM from the template text.
2. Do NOT expand abbreviations (e.g. keep "IB 2.3" not \
"Investigator's Brochure section 2.3").
3. Do NOT infer sources from industry conventions or your own knowledge.
4. Do NOT add sources that are not literally written in the section text.
5. If no source is explicitly named, return an empty list.
6. Include a "notes" field with a brief factual statement about what \
the template says (start with "Template states:").

Return a JSON object with key "sections" containing a list of objects:
  - "section_id": matching the input section_id
  - "required_sources": list of verbatim source strings
  - "notes": factual note about what the template states\
"""


def parse_template(
    template_path: Path,
    config: Config,
    llm: LLMClient,
) -> list[TemplateSection]:
    """Parse a template file (.txt or .docx) into TemplateSection objects.

    Supports .docx with mapping table parsing and IGNORE section handling.
    Checks cache first; only calls the API if the template has changed.
    """
    text, tables = _read_template_content(template_path)
    text_hash = _template_hash(template_path)
    cache_dir = ensure_dir(config.intermediate_dir)
    cache_path = cache_dir / f"parsed_template_{text_hash}.json"

    cached = _load_cache(cache_path)
    if cached is not None:
        logger.info("Using cached template parse (%d sections)", len(cached))
        return cached

    logger.info("Parsing template: %s", template_path.name)

    # --- Call 1: identify sections ---
    section_data = llm.call_json(
        system_prompt=SECTION_ID_SYSTEM,
        user_prompt=text,
        label="template_sections",
    )
    raw_sections = section_data.get("sections", [])
    logger.info("Identified %d template sections", len(raw_sections))

    # Build user prompt for source extraction with section bodies
    sections_for_source = json.dumps(
        [
            {
                "section_id": s.get("section_id", ""),
                "title": s.get("title", ""),
                "body": s.get("body", ""),
            }
            for s in raw_sections
        ],
        indent=2,
    )

    # --- Call 2: extract sources ---
    source_data = llm.call_json(
        system_prompt=SOURCE_EXTRACT_SYSTEM,
        user_prompt=sections_for_source,
        label="template_sources",
    )
    source_map: dict[str, dict] = {}
    for s in source_data.get("sections", []):
        source_map[s.get("section_id", "")] = s

    # --- Merge and validate ---
    sections: list[TemplateSection] = []
    for raw in raw_sections:
        sid = raw.get("section_id", "")
        src_info = source_map.get(sid, {})
        raw_sources = src_info.get("required_sources", [])

        # Post-extraction validation: every source must appear verbatim
        validated_sources: list[str] = []
        for src in raw_sources:
            if src in text:
                validated_sources.append(src)
            else:
                logger.warning(
                    "Dropped non-verbatim source '%s' from section %s", src, sid
                )

        section = TemplateSection(
            section_id=sid,
            title=raw.get("title", ""),
            body=raw.get("body", ""),
            required_sources=validated_sources,
            notes=src_info.get("notes", ""),
        )
        sections.append(section)

    # --- Apply mapping table sources (from .docx tables) ---
    mapping_entries = _parse_mapping_table(tables)
    if mapping_entries:
        logger.info("Found %d mapping table entries", len(mapping_entries))
        _apply_mapping_table(sections, mapping_entries)

    # --- Mark and filter IGNORE sections ---
    _mark_ignore_sections(sections)
    sections = [s for s in sections if not s.ignore]

    logger.info("Template parsing complete: %d sections, sources validated", len(sections))
    _save_cache(cache_path, sections, text_hash)
    return sections
