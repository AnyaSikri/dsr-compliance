"""Tests for template parser enhancements (.docx, mapping table, IGNORE)."""

from __future__ import annotations

from pathlib import Path

from src.models import MappingTableEntry, TemplateSection
from src.template_parser import (
    _mark_ignore_sections,
    _parse_mapping_table,
    _read_template_content,
    _split_source_refs,
)


class TestReadTemplateContent:
    def test_txt_file_returns_text_and_empty_tables(self, tmp_path: Path) -> None:
        txt = tmp_path / "template.txt"
        txt.write_text("Section 1\nContent here")
        text, tables = _read_template_content(txt)
        assert "Section 1" in text
        assert tables == []

    def test_docx_file_extracts_paragraphs(self, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph("1.1 Introduction")
        doc.add_paragraph("Some body text here")
        docx_path = tmp_path / "template.docx"
        doc.save(str(docx_path))

        text, tables = _read_template_content(docx_path)
        assert "1.1 Introduction" in text
        assert "Some body text here" in text

    def test_docx_extracts_tables(self, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "DSR Section"
        table.cell(0, 1).text = "Content"
        table.cell(0, 2).text = "Source Documents"
        table.cell(1, 0).text = "1.2.1"
        table.cell(1, 1).text = "Drug Background"
        table.cell(1, 2).text = "IB 2.3"
        docx_path = tmp_path / "template.docx"
        doc.save(str(docx_path))

        text, tables = _read_template_content(docx_path)
        assert len(tables) == 1
        assert tables[0][0][0] == "DSR Section"
        assert tables[0][1][2] == "IB 2.3"


class TestSplitSourceRefs:
    def test_split_on_or(self) -> None:
        refs = _split_source_refs("IB 6.1 OR PBRER 1.3")
        assert refs == ["IB 6.1", "PBRER 1.3"]

    def test_split_on_comma(self) -> None:
        refs = _split_source_refs("IB 2.3, IB 3.1, PBRER 5.1.2")
        assert len(refs) == 3

    def test_split_on_semicolon(self) -> None:
        refs = _split_source_refs("IB 2.3; PBRER 1.1")
        assert len(refs) == 2

    def test_single_ref(self) -> None:
        refs = _split_source_refs("IB 4.3.3")
        assert refs == ["IB 4.3.3"]

    def test_empty_string(self) -> None:
        refs = _split_source_refs("")
        assert refs == []


class TestParseMappingTable:
    def test_identifies_mapping_table(self) -> None:
        tables = [
            [
                ["DSR Section", "Content", "Source Documents"],
                ["1.2.1", "Drug Background", "IB 2.3"],
                ["1.2.1.1", "Therapeutic Indications", "IB 6.1 OR PBRER 1.3"],
            ]
        ]
        entries = _parse_mapping_table(tables)
        assert len(entries) == 2
        assert entries[0].dsr_section_id == "1.2.1"
        assert entries[0].source_refs == ["IB 2.3"]
        assert entries[1].dsr_section_id == "1.2.1.1"
        assert "IB 6.1" in entries[1].source_refs
        assert "PBRER 1.3" in entries[1].source_refs

    def test_no_tables_returns_empty(self) -> None:
        entries = _parse_mapping_table([])
        assert entries == []

    def test_table_without_keywords_skipped(self) -> None:
        tables = [
            [
                ["Name", "Age", "City"],
                ["Alice", "30", "NYC"],
            ]
        ]
        entries = _parse_mapping_table(tables)
        assert entries == []

    def test_empty_rows_skipped(self) -> None:
        tables = [
            [
                ["DSR Section", "Content", "Source Documents"],
                ["", "", ""],
                ["1.1", "Introduction", "IB 2.3"],
            ]
        ]
        entries = _parse_mapping_table(tables)
        assert len(entries) == 1


class TestIgnoreSections:
    def test_ignore_section_marks_remaining(self) -> None:
        sections = [
            TemplateSection(section_id="1", title="Introduction"),
            TemplateSection(section_id="2", title="Background"),
            TemplateSection(section_id="99", title="IGNORE: Previous Template Version"),
            TemplateSection(section_id="100", title="Old Section"),
        ]
        _mark_ignore_sections(sections)
        assert sections[0].ignore is False
        assert sections[1].ignore is False
        assert sections[2].ignore is True
        assert sections[3].ignore is True

    def test_no_ignore_marker(self) -> None:
        sections = [
            TemplateSection(section_id="1", title="Introduction"),
            TemplateSection(section_id="2", title="Background"),
        ]
        _mark_ignore_sections(sections)
        assert all(not s.ignore for s in sections)

    def test_previous_template_version_in_body(self) -> None:
        sections = [
            TemplateSection(section_id="1", title="Good Section"),
            TemplateSection(
                section_id="2",
                title="Some Title",
                body="Previous Template Version content here",
            ),
        ]
        _mark_ignore_sections(sections)
        assert sections[0].ignore is False
        assert sections[1].ignore is True
